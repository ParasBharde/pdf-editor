"""
DOCX Conversion Service
Handles PDF to DOCX conversion and DOCX manipulation
"""
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from typing import List, Optional, Dict
from app.services.redaction_patterns import RedactionPatterns


class DOCXService:
    """
    Service for DOCX conversion and manipulation
    """

    @staticmethod
    def pdf_to_docx(pdf_bytes: bytes) -> bytes:
        """
        Convert PDF to DOCX

        Args:
            pdf_bytes: PDF file as bytes

        Returns:
            DOCX file as bytes
        """
        # Create temporary streams
        pdf_stream = io.BytesIO(pdf_bytes)
        docx_stream = io.BytesIO()

        # Convert PDF to DOCX
        cv = Converter(pdf_stream)
        cv.convert(docx_stream)
        cv.close()

        docx_stream.seek(0)
        return docx_stream.read()

    @staticmethod
    def redact_docx(docx_bytes: bytes, redaction_types: List[str]) -> bytes:
        """
        Redact sensitive information from DOCX

        Args:
            docx_bytes: DOCX file as bytes
            redaction_types: List of data types to redact

        Returns:
            Redacted DOCX as bytes
        """
        # Load DOCX
        doc = Document(io.BytesIO(docx_bytes))

        # Get all text to find what needs to be redacted
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        sensitive_data = RedactionPatterns.get_redaction_items(full_text, redaction_types)

        # Collect all terms to redact
        terms_to_redact = set()
        for data_type, items in sensitive_data.items():
            terms_to_redact.update(items)

        # Redact in paragraphs
        for paragraph in doc.paragraphs:
            for term in terms_to_redact:
                if term in paragraph.text:
                    # Replace with [REDACTED]
                    paragraph.text = paragraph.text.replace(term, '[REDACTED]')

        # Redact in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for term in terms_to_redact:
                            if term in paragraph.text:
                                paragraph.text = paragraph.text.replace(term, '[REDACTED]')

        # Save to bytes
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        return output_stream.read()

    @staticmethod
    def add_header_footer_docx(docx_bytes: bytes, header_config: Optional[Dict] = None,
                               footer_config: Optional[Dict] = None,
                               logo_bytes: Optional[bytes] = None) -> bytes:
        """
        Add header and footer to DOCX document

        Args:
            docx_bytes: DOCX file as bytes
            header_config: Header configuration
            footer_config: Footer configuration
            logo_bytes: Logo image bytes (PNG, JPG)

        Returns:
            Modified DOCX as bytes
        """
        doc = Document(io.BytesIO(docx_bytes))

        # Add header
        if header_config or logo_bytes:
            section = doc.sections[0]
            header = section.header

            # Clear existing header
            for paragraph in header.paragraphs:
                paragraph.clear()

            # Add logo if provided
            if logo_bytes:
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run = header_para.add_run()

                try:
                    logo_stream = io.BytesIO(logo_bytes)
                    run.add_picture(logo_stream, width=Inches(1.5))
                except Exception as e:
                    print(f"Warning: Could not add logo to header: {e}")

            # Add header text if provided
            if header_config and header_config.get('text'):
                header_para = header.add_paragraph(header_config['text'])
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                header_para.runs[0].font.size = Pt(header_config.get('font_size', 10))

        # Add footer
        if footer_config:
            section = doc.sections[0]
            footer = section.footer

            # Clear existing footer
            for paragraph in footer.paragraphs:
                paragraph.clear()

            # Add footer text
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = footer_config.get('text', '')

            # Format footer
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT
            }
            footer_para.alignment = align_map.get(footer_config.get('align', 'center'), WD_ALIGN_PARAGRAPH.CENTER)
            footer_para.runs[0].font.size = Pt(footer_config.get('font_size', 9))
            footer_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # Save to bytes
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        return output_stream.read()

    @staticmethod
    def process_docx_full(pdf_bytes: bytes, redaction_types: List[str],
                         header_config: Optional[Dict] = None,
                         footer_config: Optional[Dict] = None,
                         logo_bytes: Optional[bytes] = None) -> bytes:
        """
        Full DOCX processing pipeline: PDF -> DOCX -> Redact -> Add Header/Footer

        Args:
            pdf_bytes: Original PDF as bytes
            redaction_types: List of data types to redact
            header_config: Header configuration
            footer_config: Footer configuration
            logo_bytes: Logo image bytes

        Returns:
            Processed DOCX as bytes
        """
        # Step 1: Convert PDF to DOCX
        docx_bytes = DOCXService.pdf_to_docx(pdf_bytes)

        # Step 2: Redact sensitive information
        if redaction_types:
            docx_bytes = DOCXService.redact_docx(docx_bytes, redaction_types)

        # Step 3: Add header and footer
        if header_config or footer_config or logo_bytes:
            docx_bytes = DOCXService.add_header_footer_docx(
                docx_bytes, header_config, footer_config, logo_bytes
            )

        return docx_bytes


class DOCXValidator:
    """
    Validator for DOCX operations
    """

    @staticmethod
    def validate_image(image_bytes: bytes) -> bool:
        """
        Validate image file

        Args:
            image_bytes: Image as bytes

        Returns:
            True if valid image
        """
        try:
            from PIL import Image
            Image.open(io.BytesIO(image_bytes))
            return True
        except Exception:
            return False
